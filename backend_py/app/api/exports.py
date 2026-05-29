from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response

from ..auth import get_current_user_id


router = APIRouter()


# ============= Markdown Generation Helper =============


def generate_markdown(framework_data: dict) -> str:
    """
    Convert frame data to Markdown format

    Args:
        framework_data: A dictionary containing metadata, steps, artifacts, risks, and escalation.

    Returns:
        str: Markdown formatted content
    """
    lines = []

    # ===== Header =====
    metadata = framework_data.get("metadata", {})
    title = metadata.get("title", "Framework")
    version = metadata.get("version", "1.0")

    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"**Version:** {version}")
    lines.append("")

    # Description
    description = metadata.get("description")
    if description:
        lines.append("## Description")
        lines.append("")
        lines.append(description)
        lines.append("")

    # ===== Framework Stages =====
    steps = framework_data.get("steps", [])
    if steps:
        lines.append("## Framework Stages")
        lines.append("")

        for i, step in enumerate(steps, 1):
            stage_name = step.get("name", f"Stage {i}")
            lines.append(f"### {i}. {stage_name}")
            lines.append("")

            # Stage description
            stage_desc = step.get("description")
            if stage_desc:
                lines.append(stage_desc)
                lines.append("")

            # Sub-steps
            substeps = step.get("subSteps", [])
            if substeps:
                lines.append("**Key Activities:**")
                lines.append("")
                for substep in substeps:
                    lines.append(f"- {substep}")
                lines.append("")

    # ===== Artefacts =====
    artefacts = framework_data.get("artefacts", {})
    if artefacts:
        lines.append("## Deliverables & Artefacts")
        lines.append("")

        # Input artefacts
        inputs = artefacts.get("input", [])
        if inputs:
            lines.append("### Input Artefacts")
            lines.append("")
            for item in inputs:
                name = item.get("name", "Unnamed")
                desc = item.get("description", "")
                lines.append(f"**{name}**")
                if desc:
                    lines.append(f": {desc}")
                lines.append("")

        # Output artefacts
        outputs = artefacts.get("output", [])
        if outputs:
            lines.append("### Output Artefacts")
            lines.append("")
            for item in outputs:
                name = item.get("name", "Unnamed")
                desc = item.get("description", "")
                lines.append(f"**{name}**")
                if desc:
                    lines.append(f": {desc}")
                lines.append("")

    # ===== Risks =====
    risks = framework_data.get("risks", [])
    if risks:
        lines.append("## Risk Considerations")
        lines.append("")

        for i, risk in enumerate(risks, 1):
            risk_name = risk.get("name", f"Risk {i}")
            lines.append(f"### {i}. {risk_name}")
            lines.append("")

            # Risk description
            risk_desc = risk.get("description")
            if risk_desc:
                lines.append(risk_desc)
                lines.append("")

            # Impact & Mitigation
            impact = risk.get("impact")
            mitigation = risk.get("mitigation")

            if impact:
                lines.append(f"**Impact:** {impact}")
                lines.append("")

            if mitigation:
                lines.append(f"**Mitigation:** {mitigation}")
                lines.append("")

    # ===== Escalation =====
    escalation = framework_data.get("escalation", [])
    if escalation:
        lines.append("## Escalation Path")
        lines.append("")

        for i, esc in enumerate(escalation, 1):
            trigger = esc.get("trigger", f"Trigger {i}")
            action = esc.get("action", "")

            lines.append(f"### {i}. {trigger}")
            lines.append("")

            if action:
                lines.append(action)
                lines.append("")

    # ===== Footer =====
    lines.append("---")
    lines.append("")
    lines.append(f"*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")

    return "\n".join(lines)


def generate_docx(framework_data: dict) -> bytes:
    """
    Convert the framework data to Word document format.

    Args:
        framework_data: A dictionary containing metadata, steps, artifacts, risks, and escalation.

    Returns:
        bytes: The binary content of the Word document
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
    except ImportError:
        raise ImportError("Please install python-docx: pip install python-docx")

    # Create document
    doc = Document()

    # ===== Header =====
    metadata = framework_data.get("metadata", {})
    title = metadata.get("title", "Framework")
    version = metadata.get("version", "1.0")

    # title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Version information
    version_para = doc.add_paragraph(f"Version: {version}")
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.runs[0]
    version_run.font.italic = True
    version_run.font.size = Pt(11)
    version_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # empty

    # Description
    description = metadata.get("description")
    if description:
        doc.add_heading("Description", level=1)
        doc.add_paragraph(description)
        doc.add_paragraph()

    # ===== POV (Points of View) =====
    pov = metadata.get("pov", [])
    if pov:
        doc.add_heading("Points of View", level=1)
        for i, point in enumerate(pov, 1):
            para = doc.add_paragraph(f"{i}. {point}", style="List Number")
            para.paragraph_format.left_indent = Inches(0.5)
        doc.add_paragraph()

    # ===== Framework Stages =====
    steps = framework_data.get("steps", [])
    if steps:
        doc.add_heading("Framework Stages", level=1)

        for i, step in enumerate(steps, 1):
            stage_name = step.get("name", f"Stage {i}")
            doc.add_heading(f"{i}. {stage_name}", level=2)

            # Stage description
            stage_desc = step.get("description")
            if stage_desc:
                doc.add_paragraph(stage_desc)

            # Sub-steps
            substeps = step.get("subSteps", [])
            if substeps:
                doc.add_paragraph("Key Activities:", style="Heading 3")
                for substep in substeps:
                    para = doc.add_paragraph(substep, style="List Bullet")
                    para.paragraph_format.left_indent = Inches(0.5)

            doc.add_paragraph()  # 空行

    # ===== Artefacts =====
    artefacts = framework_data.get("artefacts", {})
    if artefacts:
        doc.add_heading("Deliverables & Artefacts", level=1)

        # Input artefacts
        inputs = artefacts.get("input", [])
        if inputs:
            doc.add_heading("Input Artefacts", level=2)
            for item in inputs:
                name = item.get("name", "Unnamed")
                desc = item.get("description", "")
                para = doc.add_paragraph()
                run = para.add_run(name)
                run.bold = True
                if desc:
                    para.add_run(f": {desc}")

        # Output artefacts
        outputs = artefacts.get("output", [])
        if outputs:
            doc.add_heading("Output Artefacts", level=2)
            for item in outputs:
                name = item.get("name", "Unnamed")
                desc = item.get("description", "")
                para = doc.add_paragraph()
                run = para.add_run(name)
                run.bold = True
                if desc:
                    para.add_run(f": {desc}")

        # Additional artefacts
        additional = artefacts.get("additional", [])
        if additional:
            doc.add_heading("Additional Artefacts", level=2)
            for item in additional:
                name = item.get("name", "Unnamed")
                desc = item.get("description", "")
                para = doc.add_paragraph()
                run = para.add_run(name)
                run.bold = True
                if desc:
                    para.add_run(f": {desc}")

        doc.add_paragraph()

    # ===== Risks =====
    risks = framework_data.get("risks", [])
    if risks:
        doc.add_heading("Risk Considerations", level=1)

        for i, risk in enumerate(risks, 1):
            risk_name = risk.get("name", f"Risk {i}")
            doc.add_heading(f"{i}. {risk_name}", level=2)

            # Risk description
            risk_desc = risk.get("description")
            if risk_desc:
                doc.add_paragraph(risk_desc)

            # Impact & Mitigation
            impact = risk.get("impact")
            mitigation = risk.get("mitigation")

            if impact:
                para = doc.add_paragraph()
                run = para.add_run("Impact: ")
                run.bold = True
                para.add_run(impact)

            if mitigation:
                para = doc.add_paragraph()
                run = para.add_run("Mitigation: ")
                run.bold = True
                para.add_run(mitigation)

            doc.add_paragraph()

    # ===== Escalation =====
    escalation = framework_data.get("escalation", [])
    if escalation:
        doc.add_heading("Escalation Path", level=1)

        for i, esc in enumerate(escalation, 1):
            trigger = esc.get("trigger", f"Trigger {i}")
            action = esc.get("action", "")

            doc.add_heading(f"{i}. {trigger}", level=2)
            if action:
                doc.add_paragraph(action)

        doc.add_paragraph()

    # ===== Footer =====
    doc.add_paragraph()
    footer_para = doc.add_paragraph(
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


# Export endpoint
@router.post("/export-markdown")
async def export_markdown_from_data(
    framework_data: dict, current_user_id: str = Depends(get_current_user_id)
):
    """
    Receive complete framework data, generate and return Markdown files.

    Request Body:
    {
      "id": "framework-xxx",
      "metadata": {...},
      "steps": [...],
      "artefacts": {...},
      "risks": [...],
      "escalation": [...]
    }
    """
    try:
        # Generate Markdown content
        markdown_content = generate_markdown(framework_data)

        # Generate filename
        title = framework_data.get("metadata", {}).get("title", "framework")
        # Clean up filenames (remove special characters)
        safe_title = "".join(
            c if c.isalnum() or c in (" ", "-", "_") else "_" for c in title
        )
        filename = f"{safe_title.replace(' ', '_')}.md"

        # Return file
        return Response(
            content=markdown_content,
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        print(" Export Error:")
        print(traceback.format_exc())

        raise HTTPException(
            status_code=500, detail=f"Failed to export markdown: {str(e)}"
        )


@router.post("/export-docx")
async def export_docx_from_data(
    framework_data: dict, current_user_id: str = Depends(get_current_user_id)
):
    """
    Receive complete framework data, generate and return a Word document.

    Request Body:
    {
      "id": "framework-xxx",
      "metadata": {...},
      "steps": [...],
      "artefacts": {...},
      "risks": [...],
      "escalation": [...]
    }
    """
    try:
        # Generate Word document content
        docx_content = generate_docx(framework_data)

        # Generate filename
        title = framework_data.get("metadata", {}).get("title", "framework")
        # Clean up filenames (remove special characters)
        safe_title = "".join(
            c if c.isalnum() or c in (" ", "-", "_") else "_" for c in title
        )
        filename = f"{safe_title.replace(' ', '_')}.docx"

        # Return file
        return Response(
            content=docx_content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        print("❌ Export DOCX Error:")
        print(traceback.format_exc())

        raise HTTPException(
            status_code=500, detail=f"Failed to export Word document: {str(e)}"
        )
