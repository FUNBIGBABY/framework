import json
import os
import random
import sys
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from fastapi import HTTPException
from nanoid import generate
from pydantic import BaseModel
from sqlalchemy.orm import Session

from ..models import FRAMEWORK_GROUPS, Framework
from ..services.llm import (
    LLMProviderError,
    get_llm_provider,
    sanitize_model_for_provider,
)


backend_root = Path(__file__).resolve().parents[2]
if str(backend_root) not in sys.path:
    sys.path.append(str(backend_root))


class TextGenerateRequest(BaseModel):
    text: str
    use_global_llm: bool = True
    model: Optional[str] = None


class GenerateResponse(BaseModel):
    success: bool
    framework_id: Optional[str] = None
    framework: Optional[dict] = None
    frameworks: Optional[List[dict]] = None  # Multiple framework
    # Local LLM
    metadata: Optional[dict] = None
    error: Optional[str] = None


class RegenerateRequest(BaseModel):
    framework: dict
    use_local: bool = False


class FrameworkListResponse(BaseModel):
    """Frame list response"""

    id: str
    title: str
    version: str
    family: str
    confidence: float
    created_at: datetime
    updated_at: datetime

    # Simplified content preview (for card display)
    preview_artefacts: List[dict]  # Maximum of 3 artifacts


class FrameworkDetailResponse(BaseModel):
    """Framework details response"""

    id: str
    title: str
    version: str
    family: str
    confidence: float
    creator_id: str
    metadata: dict
    steps: List[dict]
    artefacts: dict
    risks: List[dict]
    escalation: List[dict]
    created_at: datetime
    updated_at: datetime


def calculate_mock_confidence() -> float:
    """
    Generate mock confidence scores (60-95)
    In the future, AI can be used to calculate the true confidence level.
    """
    return round(random.uniform(60, 95), 1)


def _env_enabled(name: str, default: bool = False) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def _require_legacy_llm_enabled() -> None:
    if not _env_enabled("ENABLE_LEGACY_LLM", False):
        raise HTTPException(
            status_code=503,
            detail=(
                "Legacy local/Ollama LLM path is disabled. Set "
                "ENABLE_LEGACY_LLM=true only for explicit legacy testing."
            ),
        )


def read_file_content(file_path: str, filename: str) -> str:
    """智能读取文件内容，根据文件类型使用不同的方法"""
    file_ext = filename.lower().split(".")[-1] if "." in filename else ""

    # 处理 .docx 文件
    if file_ext == "docx":
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n".join(paragraphs)

            if not content.strip():
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        if row_text.strip():
                            tables_text.append(row_text)
                content = "\n".join(tables_text)

            return content
        except Exception as e:
            print(f"Warning: Failed to read .docx file {filename}: {e}")

    # 处理文本文件
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1", "cp1252"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
                if content and not content.startswith("PK"):
                    return content
        except (UnicodeDecodeError, UnicodeError):
            continue

    return f"[Unable to read file: {filename}]"


def ensure_family_in_framework(framework: dict) -> str:
    family = framework.get("family") or framework.get("category")

    if family and family in FRAMEWORK_GROUPS:
        return family

    # If the AI ​​does not return a value or returns an invalid value, try to make an intelligent inference based on the title.
    title = framework.get("title", "").lower()
    # Expand the keyword list and sort by priority.
    # Technology & AI
    if any(
        word in title
        for word in [
            "ai",
            "artificial intelligence",
            "machine learning",
            "ml",
            "tech",
            "software",
            "system",
            "platform",
            "data",
            "algorithm",
            "digital",
            "cloud",
            "api",
            "code",
            "programming",
        ]
    ):
        return "Technology"

    # Healthcare & Wellbeing   add wellbeing, wellness
    elif any(
        word in title
        for word in [
            "health",
            "medical",
            "patient",
            "hospital",
            "clinical",
            "wellbeing",
            "wellness",
            "healthcare",
            "care",
            "medicine",
            "diagnosis",
            "treatment",
            "therapy",
            "pharmaceutical",
        ]
    ):
        return "Healthcare"

    # Legal & Compliance
    elif any(
        word in title
        for word in [
            "legal",
            "law",
            "compliance",
            "regulation",
            "regulatory",
            "audit",
            "governance",
            "policy",
            "risk management",
            "gdpr",
            "privacy",
            "data protection",
            "contract",
        ]
    ):
        return "Legal"

    # Financial
    elif any(
        word in title
        for word in [
            "finance",
            "financial",
            "bank",
            "invest",
            "investment",
            "accounting",
            "treasury",
            "payment",
            "trading",
            "fund",
            "capital",
            "credit",
            "loan",
            "insurance",
        ]
    ):
        return "Financial"

    # Education & Training
    elif any(
        word in title
        for word in [
            "education",
            "training",
            "learning",
            "course",
            "curriculum",
            "teaching",
            "student",
            "academic",
            "school",
            "university",
            "certification",
            "workshop",
        ]
    ):
        return "Education"

    # Marketing & Brand
    elif any(
        word in title
        for word in [
            "marketing",
            "brand",
            "campaign",
            "advertising",
            "promotion",
            "social media",
            "seo",
            "content marketing",
            "pr",
            "communication",
            "outreach",
        ]
    ):
        return "Marketing"

    # Operations & Process
    elif any(
        word in title
        for word in [
            "operation",
            "process",
            "workflow",
            "supply chain",
            "logistics",
            "manufacturing",
            "production",
            "delivery",
            "optimization",
            "efficiency",
        ]
    ):
        return "Operations"

    # Human Resources
    elif any(
        word in title
        for word in [
            "hr",
            "human resource",
            "recruit",
            "employee",
            "talent",
            "hiring",
            "onboarding",
            "performance",
            "compensation",
            "benefits",
            "workforce",
        ]
    ):
        return "Human Resources"

    # Sales & Business Development
    elif any(
        word in title
        for word in [
            "sales",
            "sell",
            "selling",
            "revenue",
            "customer",
            "business development",
            "account management",
            "crm",
            "pipeline",
            "deal",
        ]
    ):
        return "Sales"

    # Design & UX
    elif any(
        word in title
        for word in [
            "design",
            "ux",
            "ui",
            "user experience",
            "interface",
            "product design",
            "visual",
            "creative",
            "prototype",
            "wireframe",
            "mockup",
        ]
    ):
        return "Design"

    # Research & Analysis
    elif any(
        word in title
        for word in [
            "research",
            "study",
            "analysis",
            "investigation",
            "survey",
            "questionnaire",
            "data collection",
            "findings",
            "methodology",
            "hypothesis",
        ]
    ):
        return "Research"

    # Strategy & Planning
    elif any(
        word in title
        for word in [
            "strategy",
            "strategic",
            "planning",
            "roadmap",
            "business plan",
            "vision",
            "mission",
            "objectives",
            "goals",
            "initiative",
        ]
    ):
        return "Strategy"

    # Project Management
    elif any(
        word in title
        for word in [
            "project",
            "program",
            "delivery",
            "implementation",
            "milestone",
            "sprint",
            "agile",
            "scrum",
            "waterfall",
            "gantt",
            "timeline",
        ]
    ):
        return "Project Management"

    # Default fallback
    else:
        return "Other"


def build_mock_framework(metadata: dict) -> dict:
    title = (
        metadata.get("title")
        or metadata.get("subject")
        or metadata.get("document_title")
        or "Untitled Framework"
    )
    keywords = (
        metadata.get("keywords") if isinstance(metadata.get("keywords"), list) else []
    )
    sections = (
        metadata.get("sections") if isinstance(metadata.get("sections"), list) else []
    )

    steps = []
    for index, section in enumerate(sections[:4], 1):
        if isinstance(section, dict):
            name = section.get("title") or f"Stage {index}"
            description = section.get("content") or section.get("preview") or ""
        else:
            name = f"Stage {index}"
            description = str(section)
        steps.append(
            {
                "id": f"step-{index}",
                "name": str(name)[:120],
                "description": str(description)[:500],
                "subSteps": [
                    "Review available context",
                    "Identify decisions and owners",
                    "Document next actions",
                ],
            }
        )

    if not steps:
        steps = [
            {
                "id": "step-1",
                "name": "Discover",
                "description": "Clarify goals, constraints, stakeholders, and source evidence.",
                "subSteps": [
                    "Collect context",
                    "Map stakeholders",
                    "Confirm objectives",
                ],
            },
            {
                "id": "step-2",
                "name": "Design",
                "description": "Translate the source material into a structured working approach.",
                "subSteps": [
                    "Define stages",
                    "Assign ownership",
                    "Set review criteria",
                ],
            },
            {
                "id": "step-3",
                "name": "Operate",
                "description": "Run the framework, monitor outcomes, and improve it over time.",
                "subSteps": ["Execute workflow", "Track risks", "Review outcomes"],
            },
        ]

    framework = {
        "metadata": {
            "title": str(title).strip() or "Untitled Framework",
            "version": "1.0.0",
            "description": "Mock framework generated because an explicit mock path was requested.",
            "tags": [str(keyword) for keyword in keywords[:8]],
        },
        "steps": steps,
        "artefacts": {
            "input": [],
            "output": [
                {
                    "name": "Framework Brief",
                    "description": "Summary of the framework structure and key decisions.",
                }
            ],
            "additional": [],
        },
        "risks": [],
        "escalation": [],
        "pov": "general",
        "confidence": calculate_mock_confidence(),
    }
    framework["family"] = ensure_family_in_framework(
        {"title": framework["metadata"]["title"], "family": metadata.get("family")}
    )
    return framework


def process_with_local_llm(input_data: str, is_file: bool = False) -> dict:
    """
    Step 1: Extract metadata using the legacy local/Ollama path.

    Args:
        input_data: File path or text content
        is_file: Is it a file path?

    Returns:
        metadata: Extracted structured metadata
    """
    try:
        _require_legacy_llm_enabled()
        from llm_local import extract_seed

        print(
            f" Step 1: Processing {'file' if is_file else 'text'} with legacy Local LLM..."
        )

        metadata = extract_seed(input_data=input_data)

        return metadata

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Local LLM processing failed: {str(e)}"
        )


def process_with_global_llm(
    metadata: dict, model: Optional[str] = None, use_mock: bool = False
) -> dict:
    """
    Step 2: Generate the framework using the configured LLM Provider.

    Note: This will allow AI to automatically assign the family field.
    """
    try:
        if use_mock:
            print("  Using mock framework generation because mock mode was requested")
            framework = build_mock_framework(metadata)
        else:
            print(" Calling configured LLM Provider")
            provider = get_llm_provider()
            sanitized_model = sanitize_model_for_provider(
                model,
                provider_name=getattr(provider, "name", None),
            )
            framework = provider.generate_json(
                [
                    {
                        "role": "system",
                        "content": (
                            "You generate structured business frameworks from extracted "
                            "document metadata. Return only valid JSON with metadata, "
                            "steps, artefacts, risks, escalation, family, confidence, and pov fields."
                        ),
                    },
                    {
                        "role": "user",
                        "content": (
                            "Generate a framework from this metadata JSON:\n"
                            f"{json.dumps(metadata, ensure_ascii=False, indent=2)}"
                        ),
                    },
                ],
                model=sanitized_model,
                temperature=0.3,
                timeout=180.0,
            )
            print(" LLM Provider call successful")

        # Ensure the family field exists.
        # framework['family'] = ensure_family_in_framework(framework)

        return framework

    except (LLMProviderError, NotImplementedError) as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        import traceback

        print(" Global LLM Error:")
        print(traceback.format_exc())

        raise HTTPException(
            status_code=500, detail=f"Global LLM processing failed: {str(e)}"
        )


def save_framework_to_db(
    framework_data: dict, metadata_dict: dict, creator_id: str, db: Session
) -> Framework:
    """
    Save the generated framework to the database.

    Args:
        framework_data: The complete framework generated by AI
        metadata_dict: Metadata extracted by Local LLM
        creator_id: Creator user ID
        db: database session

    Returns:
        Saved Framework object
    """

    # Generate Frame ID
    framework_id = f"fw_{generate(size=12)}"

    # Extracting data from each part
    metadata = framework_data.get("metadata", {})
    steps = framework_data.get("steps", [])
    artefacts = framework_data.get("artefacts", {})
    risks = framework_data.get("risks", [])
    escalation = framework_data.get("escalation", [])
    pov = framework_data.get("pov")
    family = framework_data.get("family", "Other")
    confidence = float(framework_data.get("confidence", 0))

    # Get basic information
    title = metadata.get("title") or framework_data.get("title", "Untitled Framework")
    version = metadata.get("version", "1.0.0")
    # family = ensure_family_in_framework(framework_data)
    # confidence = calculate_mock_confidence()

    # new
    family = framework_data.get("family", "Other")
    confidence = float(framework_data.get("confidence", 0))
    pov = framework_data.get("pov", None)

    # Create database records
    db_framework = Framework(
        id=framework_id,
        title=title,
        version=version,
        creator_id=creator_id,
        metadata_json=json.dumps(metadata, ensure_ascii=False),
        steps_json=json.dumps(steps, ensure_ascii=False),
        artefacts_json=json.dumps(artefacts, ensure_ascii=False),
        risks_json=json.dumps(risks, ensure_ascii=False),
        escalation_json=json.dumps(escalation, ensure_ascii=False),
        raw_framework_json=json.dumps(framework_data, ensure_ascii=False),
        raw_metadata_json=json.dumps(metadata_dict, ensure_ascii=False),
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
        pov=pov,
        family=family,
        confidence=confidence,
    )

    db.add(db_framework)
    db.commit()
    db.refresh(db_framework)

    return db_framework
