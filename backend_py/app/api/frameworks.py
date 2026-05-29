from fastapi import (
    APIRouter,
    UploadFile,
    File,
    HTTPException,
    BackgroundTasks,
    Depends,
)
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Optional, List
from sqlalchemy.orm import Session
import json
import tempfile
import os
import random
from pathlib import Path
from datetime import datetime
from nanoid import generate
from pydantic import BaseModel
from typing import List

# Database
from ..db import get_db
from ..models import Framework, FRAMEWORK_GROUPS, SyncedVectorItem
from ..auth import get_current_user_id

# LLM
import sys

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

try:
    from llm_local import extract_seed
    from llm_global import (
        build_mock_framework,
        call_openai_framework,
        resolve_api_settings,
    )
except ImportError as e:
    print(f"Warning: Could not import LLM modules: {e}")
    print("Make sure llm_local.py and llm_global.py are in the correct location")


router = APIRouter(prefix="/api/frameworks", tags=["frameworks"])


# ============= Request/Response Models =============


class TextGenerateRequest(BaseModel):
    text: str
    use_global_llm: bool = True
    model: str = "gpt-4o"


class GenerateResponse(BaseModel):
    success: bool
    framework_id: Optional[str] = None
    framework: Optional[dict] = None
    frameworks: Optional[List[dict]] = None  # Multiple framework
    # Local LLM
    metadata: Optional[dict] = None
    error: Optional[str] = None


class RegenerateRequest(BaseModel):
    framework: dict
    use_local: bool = False


class FrameworkListResponse(BaseModel):
    """Frame list response"""

    id: str
    title: str
    version: str
    family: str
    confidence: float
    created_at: datetime
    updated_at: datetime

    # Simplified content preview (for card display)
    preview_artefacts: List[dict]  # Maximum of 3 artifacts


class FrameworkDetailResponse(BaseModel):
    """Framework details response"""

    id: str
    title: str
    version: str
    family: str
    confidence: float
    creator_id: str
    metadata: dict
    steps: List[dict]
    artefacts: dict
    risks: List[dict]
    escalation: List[dict]
    created_at: datetime
    updated_at: datetime


class SyncLibraryRequest(BaseModel):
    project_id: Optional[str] = None
    api_key: Optional[str] = None
    id_token: Optional[str] = None
    vector_store_id: Optional[str] = None
    limit: int = 1000
    include_organization: bool = True


class EventLogRequest(BaseModel):
    type: str
    framework_id: Optional[str] = None
    tenant_id: Optional[str] = None
    user_id: Optional[str] = None
    payload: Optional[dict] = None


class PushFrameworkRequest(BaseModel):
    framework: dict
    vector_store_id: Optional[str] = None


# ============= Helper Functions =============


def calculate_mock_confidence() -> float:
    """
    Generate mock confidence scores (60-95)
    In the future, AI can be used to calculate the true confidence level.
    """
    return round(random.uniform(60, 95), 1)


def _openai_client(use_vector_store_key=False):
    from openai import OpenAI

    if use_vector_store_key:
        # use separate key for vector store access
        k = os.getenv("OPENAI_VECTOR_STORE_API_KEY") or os.getenv("OPENAI_API_KEY")
    else:
        k = os.getenv("OPENAI_API_KEY")

    if not k:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY missing")
    return OpenAI(api_key=k)


def _vs_api(client):
    try:
        return client.vector_stores
    except AttributeError:
        return client.beta.vector_stores


def _vs_upload_json(client, vs_api, vs_id, filename, content, attributes=None):
    """
    Uploads a JSON file to OpenAI Vector Store.
    """
    import io as _io

    b = _io.BytesIO(content.encode("utf-8"))
    b.name = filename

    try:
        # create file first
        f = client.files.create(file=b, purpose="assistants")

        # try to add attributes
        try:
            return vs_api.files.create(
                vector_store_id=vs_id, file_id=f.id, attributes=attributes
            )
        except TypeError:
            # if api does not support attributes, then upload without it
            print(f"⚠️ attributes not supported, uploading without metadata")
            return vs_api.files.create(vector_store_id=vs_id, file_id=f.id)

    except Exception as e:
        print(f"Upload failed: {e}")
        # if failed, try without attributes
        b.seek(0)
        try:
            return vs_api.files.upload(vector_store_id=vs_id, file=b)
        except Exception as e2:
            print(f"Fallback upload failed: {e2}")
            raise e


def _read_api_key(provided):
    if provided:
        return provided
    v1 = os.getenv("VITE_FIREBASE_API_KEY")
    v2 = os.getenv("FIREBASE_API_KEY")
    return v1 or v2


def _read_project_id(provided):
    if provided:
        return provided
    p1 = os.getenv("FIREBASE_PROJECT_ID")
    return p1


def _ensure_id_token(api_key, provided):
    import requests as _requests

    if provided:
        return provided
    env_tok = os.getenv("FIREBASE_ID_TOKEN")
    if env_tok:
        return env_tok
    if not api_key:
        return None
    mail = os.getenv("FIREBASE_USER_EMAIL")
    pwd = os.getenv("FIREBASE_USER_PASSWORD")
    if mail and pwd:
        url = f"https://identitytoolkit.googleapis.com/v1/accounts:signInWithPassword?key={api_key}"
        try:
            r = _requests.post(
                url,
                json={"email": mail, "password": pwd, "returnSecureToken": True},
                timeout=30,
            )
            if r.status_code == 200:
                d = r.json()
                tok = d.get("idToken")
                if tok:
                    return tok
        except Exception:
            pass
    url = f"https://identitytoolkit.googleapis.com/v1/accounts:signUp?key={api_key}"
    r = _requests.post(url, json={"returnSecureToken": True}, timeout=30)
    if r.status_code != 200:
        return None
    d = r.json()
    return d.get("idToken")


def _val(v):
    if not isinstance(v, dict):
        return v
    if "stringValue" in v:
        return v["stringValue"]
    if "booleanValue" in v:
        return v["booleanValue"]
    if "integerValue" in v:
        try:
            return int(v["integerValue"])
        except Exception:
            return v["integerValue"]
    if "doubleValue" in v:
        try:
            return float(v["doubleValue"])
        except Exception:
            return v["doubleValue"]
    if "timestampValue" in v:
        return v["timestampValue"]
    if "mapValue" in v:
        fields = v.get("mapValue", {}).get("fields", {})
        return {k: _val(fields[k]) for k in fields} if fields else {}
    if "arrayValue" in v:
        arr = v.get("arrayValue", {}).get("values", [])
        return [_val(x) for x in arr] if arr else []
    return v


def read_file_content(file_path: str, filename: str) -> str:
    """智能读取文件内容，根据文件类型使用不同的方法"""
    file_ext = filename.lower().split(".")[-1] if "." in filename else ""

    # 处理 .docx 文件
    if file_ext == "docx":
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n".join(paragraphs)

            if not content.strip():
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        if row_text.strip():
                            tables_text.append(row_text)
                content = "\n".join(tables_text)

            return content
        except Exception as e:
            print(f"Warning: Failed to read .docx file {filename}: {e}")

    # 处理文本文件
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1", "cp1252"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
                if content and not content.startswith("PK"):
                    return content
        except (UnicodeDecodeError, UnicodeError):
            continue

    return f"[Unable to read file: {filename}]"


def ensure_family_in_framework(framework: dict) -> str:
    family = framework.get("family") or framework.get("category")

    if family and family in FRAMEWORK_GROUPS:
        return family

    # If the AI ​​does not return a value or returns an invalid value, try to make an intelligent inference based on the title.
    title = framework.get("title", "").lower()
    # Expand the keyword list and sort by priority.
    # Technology & AI
    if any(
        word in title
        for word in [
            "ai",
            "artificial intelligence",
            "machine learning",
            "ml",
            "tech",
            "software",
            "system",
            "platform",
            "data",
            "algorithm",
            "digital",
            "cloud",
            "api",
            "code",
            "programming",
        ]
    ):
        return "Technology"

    # Healthcare & Wellbeing   add wellbeing, wellness
    elif any(
        word in title
        for word in [
            "health",
            "medical",
            "patient",
            "hospital",
            "clinical",
            "wellbeing",
            "wellness",
            "healthcare",
            "care",
            "medicine",
            "diagnosis",
            "treatment",
            "therapy",
            "pharmaceutical",
        ]
    ):
        return "Healthcare"

    # Legal & Compliance
    elif any(
        word in title
        for word in [
            "legal",
            "law",
            "compliance",
            "regulation",
            "regulatory",
            "audit",
            "governance",
            "policy",
            "risk management",
            "gdpr",
            "privacy",
            "data protection",
            "contract",
        ]
    ):
        return "Legal"

    # Financial
    elif any(
        word in title
        for word in [
            "finance",
            "financial",
            "bank",
            "invest",
            "investment",
            "accounting",
            "treasury",
            "payment",
            "trading",
            "fund",
            "capital",
            "credit",
            "loan",
            "insurance",
        ]
    ):
        return "Financial"

    # Education & Training
    elif any(
        word in title
        for word in [
            "education",
            "training",
            "learning",
            "course",
            "curriculum",
            "teaching",
            "student",
            "academic",
            "school",
            "university",
            "certification",
            "workshop",
        ]
    ):
        return "Education"

    # Marketing & Brand
    elif any(
        word in title
        for word in [
            "marketing",
            "brand",
            "campaign",
            "advertising",
            "promotion",
            "social media",
            "seo",
            "content marketing",
            "pr",
            "communication",
            "outreach",
        ]
    ):
        return "Marketing"

    # Operations & Process
    elif any(
        word in title
        for word in [
            "operation",
            "process",
            "workflow",
            "supply chain",
            "logistics",
            "manufacturing",
            "production",
            "delivery",
            "optimization",
            "efficiency",
        ]
    ):
        return "Operations"

    # Human Resources
    elif any(
        word in title
        for word in [
            "hr",
            "human resource",
            "recruit",
            "employee",
            "talent",
            "hiring",
            "onboarding",
            "performance",
            "compensation",
            "benefits",
            "workforce",
        ]
    ):
        return "Human Resources"

    # Sales & Business Development
    elif any(
        word in title
        for word in [
            "sales",
            "sell",
            "selling",
            "revenue",
            "customer",
            "business development",
            "account management",
            "crm",
            "pipeline",
            "deal",
        ]
    ):
        return "Sales"

    # Design & UX
    elif any(
        word in title
        for word in [
            "design",
            "ux",
            "ui",
            "user experience",
            "interface",
            "product design",
            "visual",
            "creative",
            "prototype",
            "wireframe",
            "mockup",
        ]
    ):
        return "Design"

    # Research & Analysis
    elif any(
        word in title
        for word in [
            "research",
            "study",
            "analysis",
            "investigation",
            "survey",
            "questionnaire",
            "data collection",
            "findings",
            "methodology",
            "hypothesis",
        ]
    ):
        return "Research"

    # Strategy & Planning
    elif any(
        word in title
        for word in [
            "strategy",
            "strategic",
            "planning",
            "roadmap",
            "business plan",
            "vision",
            "mission",
            "objectives",
            "goals",
            "initiative",
        ]
    ):
        return "Strategy"

    # Project Management
    elif any(
        word in title
        for word in [
            "project",
            "program",
            "delivery",
            "implementation",
            "milestone",
            "sprint",
            "agile",
            "scrum",
            "waterfall",
            "gantt",
            "timeline",
        ]
    ):
        return "Project Management"

    # Default fallback
    else:
        return "Other"


def process_with_local_llm(input_data: str, is_file: bool = False) -> dict:
    """
    Step 1: Extract metadata using Local LLM (Cloud or Ollama)

    Now it is supported to read configuration from environment variables:
    - LLM_TYPE: "cloud" or "local"
    - LOCAL_LLM_URL: Cloud LLM address (e.g., http://34.87.13.228:8000/v1)

    Args:
        input_data: File path or text content
        is_file: Is it a file path?

    Returns:
        metadata: Extracted structured metadata
    """
    try:
        #  Instead of hardcoding host and model, extract_seed is read from environment variables.
        # This allows you to correctly use Cloud LLM instead of local Ollama.
        print(
            f" Step 1: Processing {'file' if is_file else 'text'} with Local LLM (Privacy Protection)..."
        )

        metadata = extract_seed(input_data=input_data)

        return metadata

    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Local LLM processing failed: {str(e)}"
        )


def process_with_global_llm(
    metadata: dict, model: str = "gpt-4o", use_mock: bool = False
) -> dict:
    """
    Step 2: Generate the framework using Global LLM (OpenAI)

    Note: This will allow AI to automatically assign the family field.
    """
    try:
        api_key, base_url = resolve_api_settings(None, None)

        if use_mock or not api_key:
            print("  Using mock framework generation (no OpenAI API key)")
            framework = build_mock_framework(metadata)
        else:
            print(f" Calling OpenAI API with model: {model}")

            #  Enhance prompts to allow AI to assign families.
            # Note: This requires modifying the prompt setting in llm_global.py.
            # Or add additional API calls here to categorize.

            framework = call_openai_framework(
                md=metadata,
                model=model,
                timeout=180,
                api_key=api_key,
                base_url=base_url,
                verbose=True,
            )
            print(" OpenAI API call successful")

        # Ensure the family field exists.
        # framework['family'] = ensure_family_in_framework(framework)

        return framework

    except Exception as e:
        import traceback

        print(" Global LLM Error:")
        print(traceback.format_exc())

        raise HTTPException(
            status_code=500, detail=f"Global LLM processing failed: {str(e)}"
        )


def save_framework_to_db(
    framework_data: dict, metadata_dict: dict, creator_id: str, db: Session
) -> Framework:
    """
    Save the generated framework to the database.

    Args:
        framework_data: The complete framework generated by AI
        metadata_dict: Metadata extracted by Local LLM
        creator_id: Creator user ID
        db: database session

    Returns:
        Saved Framework object
    """

    # Generate Frame ID
    framework_id = f"fw_{generate(size=12)}"

    # Extracting data from each part
    metadata = framework_data.get("metadata", {})
    steps = framework_data.get("steps", [])
    artefacts = framework_data.get("artefacts", {})
    risks = framework_data.get("risks", [])
    escalation = framework_data.get("escalation", [])
    pov = framework_data.get("pov")
    family = framework_data.get("family", "Other")
    confidence = float(framework_data.get("confidence", 0))

    # Get basic information
    title = metadata.get("title") or framework_data.get("title", "Untitled Framework")
    version = metadata.get("version", "1.0.0")
    # family = ensure_family_in_framework(framework_data)
    # confidence = calculate_mock_confidence()

    # new
    family = framework_data.get("family", "Other")
    confidence = float(framework_data.get("confidence", 0))
    pov = framework_data.get("pov", None)

    # Create database records
    db_framework = Framework(
        id=framework_id,
        title=title,
        version=version,
        creator_id=creator_id,
        metadata_json=json.dumps(metadata, ensure_ascii=False),
        steps_json=json.dumps(steps, ensure_ascii=False),
        artefacts_json=json.dumps(artefacts, ensure_ascii=False),
        risks_json=json.dumps(risks, ensure_ascii=False),
        escalation_json=json.dumps(escalation, ensure_ascii=False),
        raw_framework_json=json.dumps(framework_data, ensure_ascii=False),
        raw_metadata_json=json.dumps(metadata_dict, ensure_ascii=False),
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
        pov=pov,
        family=family,
        confidence=confidence,
    )

    db.add(db_framework)
    db.commit()
    db.refresh(db_framework)

    return db_framework


# ============= API Endpoints =============


@router.post("/generate-from-text", response_model=GenerateResponse)
async def generate_from_text(
    request: TextGenerateRequest,
    user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db),
):
    """
    From a text generation framework (login required)

    Call chain: Front-end text → Local LLM → Global LLM → Save to database → Return to framework
    """
    try:
        if not request.text.strip():
            raise HTTPException(status_code=400, detail="Text content is empty")

        if len(request.text) > 50000:
            raise HTTPException(
                status_code=400, detail="Text too long (max 50,000 characters)"
            )

        #  Whether to use Local LLM depends on use_global_llm.
        if not request.use_global_llm:
            #  Lock ON: Privacy Protection Mode
            print(" Step 1: Processing with Local LLM (Privacy Protection)...")
            metadata = process_with_local_llm(request.text, is_file=False)
            print(f" Local LLM completed. Extracted {len(metadata)} metadata fields")

            print(" Step 2: Processing with Global LLM...")
            framework_result = process_with_global_llm(
                metadata=metadata, model=request.model, use_mock=False
            )
            print(" Global LLM completed")
        else:
            #  Lock OFF: Quick Mode
            print(" Processing with Global LLM (Fast Mode - No Local Processing)...")

            #  1. Extract the title (first line or first 150 characters)
            lines = request.text.strip().split("\n")
            potential_title = lines[0][:150].strip() if lines else "User Content"

            #  2. Simple keyword extraction (from the title)
            # Select words with a length greater than 3, up to a maximum of 5.
            simple_keywords = [
                word.strip()
                for word in potential_title.lower().split()
                if len(word.strip()) > 3
            ][:5]

            #  3. Extract the chapter structure (first 5 paragraphs or chapters)
            # Send only the title, not the full content.
            sections = []
            current_section_lines = []

            for line in lines[:100]:  # View only the first 100 lines
                line_stripped = line.strip()
                if not line_stripped:
                    continue

                # Determine if it is a chapter title (simple rule: shorter line, or contains numbers).
                if len(line_stripped) < 100 and (
                    line_stripped[0].isdigit()
                    or line_stripped.isupper()
                    or any(
                        marker in line_stripped.lower()
                        for marker in ["step", "phase", "stage", "chapter"]
                    )
                ):
                    if current_section_lines:
                        # Save the previous section (keeping only the first 200 words as a summary).
                        content_preview = " ".join(current_section_lines)[:200]
                        sections.append(
                            {
                                "title": current_section_lines[0][:150],
                                "content": content_preview,  #  Keep only the first 200 characters
                                "level": 1,
                            }
                        )
                        current_section_lines = [line_stripped]
                    else:
                        current_section_lines = [line_stripped]
                else:
                    if (
                        len(current_section_lines) < 3
                    ):  # Each section can retain a maximum of 3 lines of preview.
                        current_section_lines.append(line_stripped)

            # Save the last section
            if current_section_lines:
                content_preview = " ".join(current_section_lines)[:200]
                sections.append(
                    {
                        "title": current_section_lines[0][:150],
                        "content": content_preview,
                        "level": 1,
                    }
                )

            # If no sections are extracted, use simple segmentation.
            if not sections:
                # Simple segmentation: 500 words per section
                text_parts = [
                    request.text[i : i + 500]
                    for i in range(0, min(len(request.text), 2500), 500)
                ]
                sections = [
                    {
                        "title": f"Section {i+1}",
                        "content": part[:200]
                        + "...",  #  Only the first 200 words of each section are retained.
                        "level": 1,
                    }
                    for i, part in enumerate(text_parts)
                ]

            #  4. Create optimized metadata (refer to the Lock ON structure).
            metadata = {
                "doc_id": f"doc-{generate(size=12)}",
                "title": potential_title,  #  Real title
                "subject": potential_title,
                "language": "en",
                "bypass_local_llm": True,
                #  Key fields
                "keywords": simple_keywords,  #  5-10 keywords
                #  sections: Contains only chapter titles and a preview of the first 200 words.
                "sections": sections[:10],  # Up to 10 sections
                #  facets: Simple topic categorization
                "facets": {
                    "main_topic": {
                        "summary": potential_title,
                        "items": [
                            {
                                "value": kw,
                                "evidence": "",
                                "location": "",
                                "confidence": 0.8,
                            }
                            for kw in simple_keywords
                        ],
                    }
                },
                #  key_values: Key-value pairs containing key information
                "key_values": [
                    {"key": "document_title", "value": potential_title},
                    {"key": "processing_mode", "value": "direct"},
                    {"key": "section_count", "value": str(len(sections))},
                ],
                #  tags: using keywords
                "tags": simple_keywords,
                # Other required fields (keep them empty)
                "triples": [],
                "questions": [],
                "risks": [],
                "actions_todo": [],
                "metrics": [],
                "tables": [],
                "figures": [],
                "extra": {
                    "processing_mode": "direct",
                    "note": "Extracted structure without full text to reduce prompt size",
                    "original_length": len(request.text),
                    "truncated": True,
                },
            }

            # Do not add raw_text or full_content!
            # ChatGPT does not require the complete original text, only the structured information.

            framework_result = process_with_global_llm(
                metadata=metadata, model=request.model, use_mock=False
            )
            print(" Global LLM completed")

        # 🔧 Modification: Supports multiple POVs / multiple frameworks in the results.
        #  Modification: Supports multiple POVs / multiple frameworks in the results.
        frameworks = framework_result.get("frameworks", [framework_result])

        print(f" Framework generation completed: {len(frameworks)} framework(s)")

        #  The generated data is returned directly without being saved to the database (it is saved to Firebase by the front end).
        return GenerateResponse(
            success=True,
            framework_id=None,  # The front-end will have an ID after creation.
            framework=frameworks[0] if frameworks else None,
            frameworks=frameworks,
            metadata=metadata,
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f" Error in generate_from_text: {str(e)}")
        import traceback

        traceback.print_exc()
        return GenerateResponse(success=False, error=str(e))


@router.post("/generate-from-file", response_model=GenerateResponse)
async def generate_from_file(
    file: UploadFile = File(...),
    use_global_llm: bool = True,
    model: str = "gpt-4o",
    user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db),
):
    """
    Generate a framework from uploaded files (login required)

    Call chain: Front-end file → Local LLM → Global LLM → Save to database → Return to framework
    """
    temp_path = None

    try:
        # Verify file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")

        # Check file type
        allowed_extensions = {".txt", ".pdf", ".doc", ".docx", ".md"}
        file_ext = Path(file.filename).suffix.lower()

        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}",
            )

        # Check file size (10MB)
        content = await file.read()
        if len(content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large (max 10MB)")

        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
            tmp.write(content)
            temp_path = tmp.name

        print(f" File saved to: {temp_path}")

        # Step 1: Extract metadata from local LLM
        print(" Step 1: Processing with Local LLM (Ollama)...")
        metadata = process_with_local_llm(temp_path, is_file=True)
        print(f" Local LLM completed. Extracted {len(metadata)} metadata fields")

        # Step 2: Global LLM Generation Framework
        print(" Step 2: Processing with Global LLM (OpenAI)...")
        framework_result = process_with_global_llm(  #  MODIFIED
            metadata=metadata, model=model, use_mock=not use_global_llm
        )
        print(" Global LLM completed. Framework generated")

        #  MODIFIED: Supports multiple POV outputs
        frameworks = framework_result.get("frameworks", [framework_result])

        #  Step 3: Save to database
        print("💾 Step 3: Saving framework(s) to database...")
        saved_ids = []  #  MODIFIED
        for fw_data in frameworks:  #  MODIFIED
            db_framework = save_framework_to_db(  #  MODIFIED
                framework_data=fw_data,  #  MODIFIED
                metadata_dict=metadata,
                creator_id=user_id,
                db=db,
            )
            saved_ids.append(db_framework.id)  #  MODIFIED
        print(f" All frameworks saved: {len(saved_ids)} total")  #  MODIFIED

        #  MODIFIED: Returns both single and multiple values ​​(backwards compatible).
        return GenerateResponse(
            success=True,
            framework_id=saved_ids[0] if saved_ids else None,  #  MODIFIED
            framework=frameworks[0] if frameworks else None,  #  MODIFIED
            frameworks=frameworks,  #  MODIFIED
            metadata=metadata,
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f" Error in generate_from_file: {str(e)}")
        import traceback

        traceback.print_exc()
        return GenerateResponse(success=False, error=str(e))
    finally:
        # Clean up temporary files
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except:
                pass


@router.post("/generate-from-files", response_model=GenerateResponse)
async def generate_from_files(
    files: List[UploadFile] = File(...),
    use_global_llm: bool = True,
    model: str = "gpt-4o",
    user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db),
):
    """
    Generate a framework from multiple files (login required)

    Multiple files will be merged.
    """
    temp_paths = []

    try:
        if not files or len(files) == 0:
            raise HTTPException(status_code=400, detail="No files provided")

        if len(files) > 10:
            raise HTTPException(status_code=400, detail="Too many files (max 10)")

        # Save all files to a temporary directory
        for file in files:
            if not file.filename:
                continue

            file_ext = Path(file.filename).suffix.lower()
            allowed_extensions = {".txt", ".pdf", ".doc", ".docx", ".md"}

            if file_ext not in allowed_extensions:
                continue

            content = await file.read()
            if len(content) > 10 * 1024 * 1024:
                raise HTTPException(
                    status_code=400, detail=f"File {file.filename} too large"
                )

            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
                tmp.write(content)
                temp_paths.append(tmp.name)

        if not temp_paths:
            raise HTTPException(status_code=400, detail="No valid files")

        print(f" Saved {len(temp_paths)} files")

        #  Whether to use Local LLM depends on use_global_llm.
        if not use_global_llm:
            #  Lock ON: Privacy Protection Mode
            print(" Step 1: Processing files with Local LLM (Privacy Protection)...")
            all_metadata = []

            for temp_path in temp_paths:
                metadata = process_with_local_llm(temp_path, is_file=True)
                all_metadata.append(metadata)

            merged_metadata = all_metadata[0] if all_metadata else {}
            if len(all_metadata) > 1:
                merged_metadata["source_count"] = len(all_metadata)
                merged_metadata["merged_from_multiple_files"] = True

            print(f" Local LLM completed. Processed {len(temp_paths)} files")

            print(" Step 2: Processing with Global LLM...")
            framework_result = process_with_global_llm(
                metadata=merged_metadata, model=model, use_mock=False
            )
            print(" Global LLM completed")
        else:
            #  Lock OFF: Quick Mode
            print(" Processing with Global LLM (Fast Mode - No Local Processing)...")

            # Read all file contents
            file_contents = []
            file_names = []
            for i, temp_path in enumerate(temp_paths):
                try:
                    # Get filename
                    original_filename = (
                        files[i].filename if i < len(files) else f"file_{i+1}"
                    )
                    file_names.append(original_filename)

                    # Use smart read function Fix .docx garbled issue
                    content = read_file_content(temp_path, original_filename)
                    if content and not content.startswith("[Unable to read"):
                        file_contents.append(content)
                    else:
                        print(f"Warning: {content}")

                except Exception as e:
                    print(f"Warning: Could not read file {temp_path}: {e}")

            # 1. Intelligent title extraction
            if len(file_contents) == 1:
                # Single file: Use the first line or the filename
                lines = file_contents[0].strip().split("\n")
                potential_title = (
                    lines[0][:150].strip()
                    if lines and len(lines[0].strip()) > 10
                    else file_names[0]
                )
            else:
                # Multiple files: Use combined descriptions
                lines = file_contents[0].strip().split("\n") if file_contents else []
                if lines and len(lines[0].strip()) > 10:
                    potential_title = lines[0][:150].strip()
                else:
                    potential_title = f"Framework from {len(file_names)} files"

            #  2. Simple Keyword Extraction
            simple_keywords = [
                word.strip()
                for word in potential_title.lower().split()
                if len(word.strip()) > 3
            ][:5]

            #  3. Extract sections (extract from all files, but keep only the first 200 words of each section)
            all_sections = []

            for idx, content in enumerate(file_contents):
                file_name = (
                    file_names[idx] if idx < len(file_names) else f"File {idx+1}"
                )
                lines = content.strip().split("\n")

                # Create sections for each file
                current_section_lines = []

                for line in lines[
                    :50
                ]:  # Only the first 50 lines of each file are viewed.
                    line_stripped = line.strip()
                    if not line_stripped:
                        continue

                    # Determine if it is a chapter title
                    if len(line_stripped) < 100 and (
                        line_stripped[0].isdigit()
                        or line_stripped.isupper()
                        or any(
                            marker in line_stripped.lower()
                            for marker in ["step", "phase", "stage", "chapter"]
                        )
                    ):
                        if current_section_lines:
                            content_preview = " ".join(current_section_lines)[:200]
                            all_sections.append(
                                {
                                    "title": f"{file_name}: {current_section_lines[0][:100]}",
                                    "content": content_preview,  #  Keep only the first 200 characters
                                    "level": 1,
                                    "source_file": file_name,
                                }
                            )
                            current_section_lines = [line_stripped]
                        else:
                            current_section_lines = [line_stripped]
                    else:
                        if len(current_section_lines) < 3:
                            current_section_lines.append(line_stripped)

                # Save the last section
                if current_section_lines:
                    content_preview = " ".join(current_section_lines)[:200]
                    all_sections.append(
                        {
                            "title": f"{file_name}: {current_section_lines[0][:100]}",
                            "content": content_preview,
                            "level": 1,
                            "source_file": file_name,
                        }
                    )

            # If no sections are extracted, create a simple section for each file.
            if not all_sections:
                all_sections = [
                    {
                        "title": file_names[i]
                        if i < len(file_names)
                        else f"File {i+1}",
                        "content": content[:200]
                        + "...",  #  Keep only the first 200 characters
                        "level": 1,
                        "source_file": file_names[i]
                        if i < len(file_names)
                        else f"File {i+1}",
                    }
                    for i, content in enumerate(file_contents)
                ]

            #  4. Create optimized metadata
            merged_metadata = {
                "doc_id": f"doc-{generate(size=12)}",
                "title": potential_title,  #  real title
                "subject": potential_title,
                "language": "en",
                "bypass_local_llm": True,
                #  keywords
                "keywords": simple_keywords,
                #  sections: Contains only chapter titles and a preview of the first 200 words.
                "sections": all_sections[:15],  # Up to 15 sections
                #  facets
                "facets": {
                    "main_topic": {
                        "summary": potential_title,
                        "items": [
                            {
                                "value": kw,
                                "evidence": "",
                                "location": "",
                                "confidence": 0.8,
                            }
                            for kw in simple_keywords
                        ],
                    },
                    "source_files": {
                        "summary": f"Content from {len(file_contents)} file(s)",
                        "items": [
                            {
                                "value": name,
                                "evidence": "",
                                "location": "",
                                "confidence": 1.0,
                            }
                            for name in file_names
                        ],
                    },
                },
                #  key_values
                "key_values": [
                    {"key": "document_title", "value": potential_title},
                    {"key": "file_count", "value": str(len(file_contents))},
                    {"key": "processing_mode", "value": "direct"},
                    {"key": "source_files", "value": ", ".join(file_names[:3])},
                ],
                #  tags
                "tags": simple_keywords,
                # Other required fields
                "source_count": len(file_contents),
                "source_files": file_names,
                "triples": [],
                "questions": [],
                "risks": [],
                "actions_todo": [],
                "metrics": [],
                "tables": [],
                "figures": [],
                "extra": {
                    "processing_mode": "direct",
                    "note": "Extracted structure without full text to reduce prompt size",
                    "file_names": file_names,
                    "total_length": sum(len(c) for c in file_contents),
                    "truncated": True,
                },
            }

            #  Key point: Do not add raw_text, full_content, or the complete combined_text!

            framework_result = process_with_global_llm(
                metadata=merged_metadata, model=model, use_mock=False
            )
            print(" Global LLM completed")

        #  MODIFIED: Supports multiple POV outputs
        frameworks = framework_result.get("frameworks", [framework_result])

        #  Step 3: Generate framework IDs (the frontend will save them to Firebase).
        print(
            "💾 Step 3: Generating framework IDs (data will be saved to Firebase by frontend)..."
        )

        saved_ids = []
        for fw_data in frameworks:
            # Generate a unique ID for each framework
            fw_id = f"fw_{generate(size=12)}"
            fw_data["id"] = fw_id  # Add ID to framework data
            saved_ids.append(fw_id)

        print(f" Generated {len(saved_ids)} framework IDs: {saved_ids}")

        #  MODIFIED: Returns both single and multiple [MODIFIED] values.
        return GenerateResponse(
            success=True,
            framework_id=saved_ids[0] if saved_ids else None,
            framework=frameworks[0] if frameworks else None,
            frameworks=frameworks,
            metadata=merged_metadata,
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f" Error: {str(e)}")
        import traceback

        traceback.print_exc()
        return GenerateResponse(success=False, error=str(e))
    finally:
        # Clean up all temporary files
        for temp_path in temp_paths:
            if os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except:
                    pass


# New feature: Retrieve all frameworks for the current user.
@router.get("/my-frameworks", response_model=List[FrameworkListResponse])
def get_my_frameworks(
    user_id: str = Depends(get_current_user_id), db: Session = Depends(get_db)
):
    """
    Get all frameworks created by the current user

    Sort in descending order of creation time List
    For the "Your Frameworks" list page
    """

    frameworks = (
        db.query(Framework)
        .filter(Framework.creator_id == user_id)
        .order_by(Framework.created_at.desc())
        .all()
    )

    result = []
    for fw in frameworks:
        # Analyze artefacts for preview
        artefacts = json.loads(fw.artefacts_json)
        additional = artefacts.get("additional", [])

        # Only the first 3 artifacts are used for card display.
        preview_artefacts = []
        if additional:
            for art in additional[:3]:
                preview_artefacts.append(
                    {
                        "name": art.get("name", ""),
                        "description": art.get("description", "")[
                            :100
                        ],  # Truncation description
                    }
                )

        result.append(
            FrameworkListResponse(
                id=fw.id,
                title=fw.title,
                version=fw.version,
                family=fw.family,
                confidence=fw.confidence,
                created_at=fw.created_at,
                updated_at=fw.updated_at,
                preview_artefacts=preview_artefacts,
            )
        )

    return result


# New feature: Retrieve frameworks by family group
@router.get("/my-frameworks/by-family")
def get_my_frameworks_by_family(
    user_id: str = Depends(get_current_user_id), db: Session = Depends(get_db)
):
    """
    Retrieve the current user's frameworks and group them by family.

    Return format:
    {
        "Financial": [framework1, framework2, ...],
        "Healthcare": [...],
        ...
    }
    """

    frameworks = (
        db.query(Framework)
        .filter(Framework.creator_id == user_id)
        .order_by(Framework.created_at.desc())
        .all()
    )

    # Grouped by family
    grouped = {}
    for fw in frameworks:
        family = fw.family or "Other"

        if family not in grouped:
            grouped[family] = []

        # Analysis of Artefacts
        artefacts = json.loads(fw.artefacts_json)
        additional = artefacts.get("additional", [])

        preview_artefacts = []
        if additional:
            for art in additional[:3]:
                preview_artefacts.append(
                    {
                        "name": art.get("name", ""),
                        "description": art.get("description", "")[:100],
                    }
                )

        grouped[family].append(
            {
                "id": fw.id,
                "title": fw.title,
                "version": fw.version,
                "family": fw.family,
                "confidence": fw.confidence,
                "created_at": fw.created_at.isoformat(),
                "updated_at": fw.updated_at.isoformat(),
                "preview_artefacts": preview_artefacts,
            }
        )

    return grouped


# New feature: Retrieve detailed information for a single framework.
@router.get("/{framework_id}", response_model=FrameworkDetailResponse)
def get_framework_detail(
    framework_id: str,
    user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db),
):
    """
    Get complete information about the framework

    Only able to access the frameworks they created
    """

    framework = (
        db.query(Framework)
        .filter(
            Framework.id == framework_id,
            Framework.creator_id
            == user_id,  # Ensure that only your own account can be accessed.
        )
        .first()
    )

    if not framework:
        raise HTTPException(
            status_code=404,
            detail="Framework not found or you don't have permission to access it",
        )

    return FrameworkDetailResponse(
        id=framework.id,
        title=framework.title,
        version=framework.version,
        family=framework.family,
        confidence=framework.confidence,
        creator_id=framework.creator_id,
        metadata=json.loads(framework.metadata_json),
        steps=json.loads(framework.steps_json),
        artefacts=json.loads(framework.artefacts_json),
        risks=json.loads(framework.risks_json),
        escalation=json.loads(framework.escalation_json),
        created_at=framework.created_at,
        updated_at=framework.updated_at,
    )


# Added: Binding information interface (/api/frameworks/{id}/binding)
@router.get("/{framework_id}/binding")
def get_framework_binding(
    framework_id: str,
    user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db),
):
    """
    Retrieve the framework's POV, family, and confidence binding information.
    (Used for displaying POV and confidence level simultaneously in a frame card or detail page on the front end)
    """
    # 🔍 Query the framework of the current user
    fw = (
        db.query(Framework)
        .filter(Framework.id == framework_id, Framework.creator_id == user_id)
        .first()
    )

    if not fw:
        raise HTTPException(
            status_code=404, detail="Framework not found or access denied"
        )

    # Trying to extract the POV (the complete content returned by AI) from raw_framework_json.
    pov_value = None
    try:
        if fw.raw_framework_json:
            raw_data = json.loads(fw.raw_framework_json)
            pov_value = raw_data.get("pov")
    except Exception:
        pov_value = None

    #  Return unified binding information
    return {
        "id": fw.id,
        "title": fw.title,
        "pov": pov_value,
        "family": fw.family,
        "confidence": fw.confidence,
        "created_at": fw.created_at,
        "updated_at": fw.updated_at,
    }


# Added: Updated framework
@router.put("/{framework_id}")
def update_framework(
    framework_id: str,
    framework_data: dict,
    user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db),
):
    """
    Update the framework (save from the Editor)

    You can only update the framework you created.
    """

    framework = (
        db.query(Framework)
        .filter(Framework.id == framework_id, Framework.creator_id == user_id)
        .first()
    )

    if not framework:
        raise HTTPException(
            status_code=404, detail="Framework not found or you don't have permission"
        )

    # Update field
    metadata = framework_data.get("metadata", {})
    framework.title = metadata.get("title", framework.title)
    framework.version = metadata.get("version", framework.version)

    # Update JSON fields
    framework.metadata_json = json.dumps(metadata, ensure_ascii=False)
    framework.steps_json = json.dumps(
        framework_data.get("steps", []), ensure_ascii=False
    )
    framework.artefacts_json = json.dumps(
        framework_data.get("artefacts", {}), ensure_ascii=False
    )
    framework.risks_json = json.dumps(
        framework_data.get("risks", []), ensure_ascii=False
    )
    framework.escalation_json = json.dumps(
        framework_data.get("escalation", []), ensure_ascii=False
    )

    framework.updated_at = datetime.utcnow()

    db.commit()
    db.refresh(framework)

    return {
        "success": True,
        "message": "Framework updated successfully",
        "framework_id": framework.id,
    }


# Added: Remove framework
@router.delete("/{framework_id}")
def delete_framework(
    framework_id: str,
    user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db),
):
    """
    remove framework

    You can only delete the framework you created.
    """

    framework = (
        db.query(Framework)
        .filter(Framework.id == framework_id, Framework.creator_id == user_id)
        .first()
    )

    if not framework:
        raise HTTPException(
            status_code=404, detail="Framework not found or you don't have permission"
        )

    db.delete(framework)
    db.commit()

    return {"success": True, "message": "Framework deleted successfully"}


# ============= Markdown Generation Helper =============


def generate_markdown(framework_data: dict) -> str:
    """
    Convert frame data to Markdown format

    Args:
        framework_data: A dictionary containing metadata, steps, artifacts, risks, and escalation.

    Returns:
        str: Markdown formatted content
    """
    lines = []

    # ===== Header =====
    metadata = framework_data.get("metadata", {})
    title = metadata.get("title", "Framework")
    version = metadata.get("version", "1.0")

    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"**Version:** {version}")
    lines.append("")

    # Description
    description = metadata.get("description")
    if description:
        lines.append("## Description")
        lines.append("")
        lines.append(description)
        lines.append("")

    # ===== Framework Stages =====
    steps = framework_data.get("steps", [])
    if steps:
        lines.append("## Framework Stages")
        lines.append("")

        for i, step in enumerate(steps, 1):
            stage_name = step.get("name", f"Stage {i}")
            lines.append(f"### {i}. {stage_name}")
            lines.append("")

            # Stage description
            stage_desc = step.get("description")
            if stage_desc:
                lines.append(stage_desc)
                lines.append("")

            # Sub-steps
            substeps = step.get("subSteps", [])
            if substeps:
                lines.append("**Key Activities:**")
                lines.append("")
                for substep in substeps:
                    lines.append(f"- {substep}")
                lines.append("")

    # ===== Artefacts =====
    artefacts = framework_data.get("artefacts", {})
    if artefacts:
        lines.append("## Deliverables & Artefacts")
        lines.append("")

        # Input artefacts
        inputs = artefacts.get("input", [])
        if inputs:
            lines.append("### Input Artefacts")
            lines.append("")
            for item in inputs:
                name = item.get("name", "Unnamed")
                desc = item.get("description", "")
                lines.append(f"**{name}**")
                if desc:
                    lines.append(f": {desc}")
                lines.append("")

        # Output artefacts
        outputs = artefacts.get("output", [])
        if outputs:
            lines.append("### Output Artefacts")
            lines.append("")
            for item in outputs:
                name = item.get("name", "Unnamed")
                desc = item.get("description", "")
                lines.append(f"**{name}**")
                if desc:
                    lines.append(f": {desc}")
                lines.append("")

    # ===== Risks =====
    risks = framework_data.get("risks", [])
    if risks:
        lines.append("## Risk Considerations")
        lines.append("")

        for i, risk in enumerate(risks, 1):
            risk_name = risk.get("name", f"Risk {i}")
            lines.append(f"### {i}. {risk_name}")
            lines.append("")

            # Risk description
            risk_desc = risk.get("description")
            if risk_desc:
                lines.append(risk_desc)
                lines.append("")

            # Impact & Mitigation
            impact = risk.get("impact")
            mitigation = risk.get("mitigation")

            if impact:
                lines.append(f"**Impact:** {impact}")
                lines.append("")

            if mitigation:
                lines.append(f"**Mitigation:** {mitigation}")
                lines.append("")

    # ===== Escalation =====
    escalation = framework_data.get("escalation", [])
    if escalation:
        lines.append("## Escalation Path")
        lines.append("")

        for i, esc in enumerate(escalation, 1):
            trigger = esc.get("trigger", f"Trigger {i}")
            action = esc.get("action", "")

            lines.append(f"### {i}. {trigger}")
            lines.append("")

            if action:
                lines.append(action)
                lines.append("")

    # ===== Footer =====
    lines.append("---")
    lines.append("")
    lines.append(f"*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")

    return "\n".join(lines)


def generate_docx(framework_data: dict) -> bytes:
    """
    Convert the framework data to Word document format.

    Args:
        framework_data: A dictionary containing metadata, steps, artifacts, risks, and escalation.

    Returns:
        bytes: The binary content of the Word document
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
    except ImportError:
        raise ImportError("Please install python-docx: pip install python-docx")

    # Create document
    doc = Document()

    # ===== Header =====
    metadata = framework_data.get("metadata", {})
    title = metadata.get("title", "Framework")
    version = metadata.get("version", "1.0")

    # title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Version information
    version_para = doc.add_paragraph(f"Version: {version}")
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.runs[0]
    version_run.font.italic = True
    version_run.font.size = Pt(11)
    version_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # empty

    # Description
    description = metadata.get("description")
    if description:
        doc.add_heading("Description", level=1)
        doc.add_paragraph(description)
        doc.add_paragraph()

    # ===== POV (Points of View) =====
    pov = metadata.get("pov", [])
    if pov:
        doc.add_heading("Points of View", level=1)
        for i, point in enumerate(pov, 1):
            para = doc.add_paragraph(f"{i}. {point}", style="List Number")
            para.paragraph_format.left_indent = Inches(0.5)
        doc.add_paragraph()

    # ===== Framework Stages =====
    steps = framework_data.get("steps", [])
    if steps:
        doc.add_heading("Framework Stages", level=1)

        for i, step in enumerate(steps, 1):
            stage_name = step.get("name", f"Stage {i}")
            doc.add_heading(f"{i}. {stage_name}", level=2)

            # Stage description
            stage_desc = step.get("description")
            if stage_desc:
                doc.add_paragraph(stage_desc)

            # Sub-steps
            substeps = step.get("subSteps", [])
            if substeps:
                doc.add_paragraph("Key Activities:", style="Heading 3")
                for substep in substeps:
                    para = doc.add_paragraph(substep, style="List Bullet")
                    para.paragraph_format.left_indent = Inches(0.5)

            doc.add_paragraph()  # 空行

    # ===== Artefacts =====
    artefacts = framework_data.get("artefacts", {})
    if artefacts:
        doc.add_heading("Deliverables & Artefacts", level=1)

        # Input artefacts
        inputs = artefacts.get("input", [])
        if inputs:
            doc.add_heading("Input Artefacts", level=2)
            for item in inputs:
                name = item.get("name", "Unnamed")
                desc = item.get("description", "")
                para = doc.add_paragraph()
                run = para.add_run(name)
                run.bold = True
                if desc:
                    para.add_run(f": {desc}")

        # Output artefacts
        outputs = artefacts.get("output", [])
        if outputs:
            doc.add_heading("Output Artefacts", level=2)
            for item in outputs:
                name = item.get("name", "Unnamed")
                desc = item.get("description", "")
                para = doc.add_paragraph()
                run = para.add_run(name)
                run.bold = True
                if desc:
                    para.add_run(f": {desc}")

        # Additional artefacts
        additional = artefacts.get("additional", [])
        if additional:
            doc.add_heading("Additional Artefacts", level=2)
            for item in additional:
                name = item.get("name", "Unnamed")
                desc = item.get("description", "")
                para = doc.add_paragraph()
                run = para.add_run(name)
                run.bold = True
                if desc:
                    para.add_run(f": {desc}")

        doc.add_paragraph()

    # ===== Risks =====
    risks = framework_data.get("risks", [])
    if risks:
        doc.add_heading("Risk Considerations", level=1)

        for i, risk in enumerate(risks, 1):
            risk_name = risk.get("name", f"Risk {i}")
            doc.add_heading(f"{i}. {risk_name}", level=2)

            # Risk description
            risk_desc = risk.get("description")
            if risk_desc:
                doc.add_paragraph(risk_desc)

            # Impact & Mitigation
            impact = risk.get("impact")
            mitigation = risk.get("mitigation")

            if impact:
                para = doc.add_paragraph()
                run = para.add_run("Impact: ")
                run.bold = True
                para.add_run(impact)

            if mitigation:
                para = doc.add_paragraph()
                run = para.add_run("Mitigation: ")
                run.bold = True
                para.add_run(mitigation)

            doc.add_paragraph()

    # ===== Escalation =====
    escalation = framework_data.get("escalation", [])
    if escalation:
        doc.add_heading("Escalation Path", level=1)

        for i, esc in enumerate(escalation, 1):
            trigger = esc.get("trigger", f"Trigger {i}")
            action = esc.get("action", "")

            doc.add_heading(f"{i}. {trigger}", level=2)
            if action:
                doc.add_paragraph(action)

        doc.add_paragraph()

    # ===== Footer =====
    doc.add_paragraph()
    footer_para = doc.add_paragraph(
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


# Export endpoint
@router.post("/export-markdown")
async def export_markdown_from_data(
    framework_data: dict, current_user_id: str = Depends(get_current_user_id)
):
    """
    Receive complete framework data, generate and return Markdown files.

    Request Body:
    {
      "id": "framework-xxx",
      "metadata": {...},
      "steps": [...],
      "artefacts": {...},
      "risks": [...],
      "escalation": [...]
    }
    """
    try:
        # Generate Markdown content
        markdown_content = generate_markdown(framework_data)

        # Generate filename
        title = framework_data.get("metadata", {}).get("title", "framework")
        # Clean up filenames (remove special characters)
        safe_title = "".join(
            c if c.isalnum() or c in (" ", "-", "_") else "_" for c in title
        )
        filename = f"{safe_title.replace(' ', '_')}.md"

        # Return file
        return Response(
            content=markdown_content,
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        print(" Export Error:")
        print(traceback.format_exc())

        raise HTTPException(
            status_code=500, detail=f"Failed to export markdown: {str(e)}"
        )


@router.post("/export-docx")
async def export_docx_from_data(
    framework_data: dict, current_user_id: str = Depends(get_current_user_id)
):
    """
    Receive complete framework data, generate and return a Word document.

    Request Body:
    {
      "id": "framework-xxx",
      "metadata": {...},
      "steps": [...],
      "artefacts": {...},
      "risks": [...],
      "escalation": [...]
    }
    """
    try:
        # Generate Word document content
        docx_content = generate_docx(framework_data)

        # Generate filename
        title = framework_data.get("metadata", {}).get("title", "framework")
        # Clean up filenames (remove special characters)
        safe_title = "".join(
            c if c.isalnum() or c in (" ", "-", "_") else "_" for c in title
        )
        filename = f"{safe_title.replace(' ', '_')}.docx"

        # Return file
        return Response(
            content=docx_content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback

        print("❌ Export DOCX Error:")
        print(traceback.format_exc())

        raise HTTPException(
            status_code=500, detail=f"Failed to export Word document: {str(e)}"
        )


@router.post("/regenerate")
async def regenerate_framework(
    request: RegenerateRequest, current_user_id: str = Depends(get_current_user_id)
):
    """
    Regenerate the framework (improved by user editing).

    Users can choose:
    1. Cloud Processing (OpenAI) - Fast, high-quality, retains all user edits
    2. Local Processing (Ollama) - Privacy priority, but details may be lost.
    """
    try:
        if request.use_local:
            # ========== Local Processing Mode ==========
            print(" Using Local Processing (Ollama)")

            # Check if Ollama is running.
            import requests

            try:
                requests.get("http://127.0.0.1:11434", timeout=2)
            except requests.exceptions.RequestException:
                raise HTTPException(
                    status_code=503,
                    detail="Ollama is not running. Please start Ollama: 'ollama serve'",
                )

            # Step 1: Convert the user-edited frame back to text (simulating reverse engineering)
            framework_text = convert_framework_to_text(request.framework)

            # Step 2: Re-extract metadata from Local LLM
            from llm_local import extract_seed_from_text, OllamaClient

            llm = OllamaClient(model="llama3.1:8b", host="http://127.0.0.1:11434")
            metadata = extract_seed_from_text(framework_text, llm=llm)

            # Step 3: Generate a new framework using Global LLM (or use a mock).
            api_key, base_url = resolve_api_settings(None, None)
            if api_key:
                from llm_global import call_openai_framework

                improved_framework = call_openai_framework(
                    md=metadata,
                    model="gpt-4o",
                    timeout=180,
                    api_key=api_key,
                    base_url=base_url,
                    verbose=True,
                )
            else:
                from llm_global import build_mock_framework

                improved_framework = build_mock_framework(metadata)

            return {
                "success": True,
                "framework": improved_framework,
                "method": "local",
                "message": "Framework regenerated using local processing",
            }

        else:
            # ========== Cloud Processing Mode (Recommended) =========
            print("☁️ Using Cloud Processing (OpenAI)")

            # Check API key
            api_key, base_url = resolve_api_settings(None, None)
            if not api_key:
                raise HTTPException(
                    status_code=400,
                    detail="OpenAI API key not configured. Please use local processing instead.",
                )

            # Send the complete framework directly to OpenAI for improvements.
            from openai import OpenAI
            import httpx
            import os

            # Clear proxy environment variables
            original_proxies = {}
            proxy_keys = [
                "HTTP_PROXY",
                "HTTPS_PROXY",
                "http_proxy",
                "https_proxy",
                "ALL_PROXY",
                "all_proxy",
                "NO_PROXY",
                "no_proxy",
            ]

            for key in proxy_keys:
                if key in os.environ:
                    original_proxies[key] = os.environ[key]
                    del os.environ[key]
            try:
                timeout = 180.0
                if base_url:
                    client = OpenAI(
                        api_key=api_key,
                        base_url=base_url,
                        timeout=timeout,
                        max_retries=2,
                    )
                else:
                    client = OpenAI(api_key=api_key, timeout=timeout, max_retries=2)

                # Build prompt
                system_prompt = (
                    "You are a framework improvement assistant. "
                    "The user has edited a framework and wants you to review and improve it. "
                    "CRITICAL: Keep ALL user modifications intact. Only fill in missing parts and suggest improvements. "
                    "Return the improved framework as valid JSON matching the original structure."
                )

                user_prompt = (
                    "Here is a framework that the user has edited:\n\n"
                    f"{json.dumps(request.framework, indent=2)}\n\n"
                    "Please:\n"
                    "1. **Keep all user modifications intact** (especially steps, risks, escalation)\n"
                    "2. Fill in missing sections if any:\n"
                    "   - Add 'trigger_context' or 'pov' if missing\n"
                    "   - Add 'inputs_required' if missing\n"
                    "   - Add 'research_required' if missing\n"
                    "   - Add 'attribution' if appropriate\n"
                    "   - Add 'quadrant' (QI/QII/QIII/QIV) if appropriate\n"
                    "3. Ensure consistency across all sections\n"
                    "4. Improve descriptions to be more specific and actionable\n"
                    "5. Return the complete improved framework as JSON\n\n"
                    "IMPORTANT: Do NOT remove or significantly change user's content. Only enhance and complete."
                )

                print(" Sending request to OpenAI...")
                response = client.chat.completions.create(
                    model="gpt-4o",
                    temperature=0.3,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                )

                result_text = response.choices[0].message.content.strip()
                print(" Received response from OpenAI")

                # Parsing JSON
                from llm_global import robust_json_loads

                improved_framework = robust_json_loads(result_text)

                return {
                    "success": True,
                    "framework": improved_framework,
                    "method": "cloud",
                    "message": "Framework regenerated using cloud processing",
                }

            finally:
                # Restore proxy settings
                for key, value in original_proxies.items():
                    os.environ[key] = value

                # Close HTTP client
                if "http_client" in locals():
                    try:
                        http_client.close()
                    except:
                        pass

    except HTTPException:
        raise
    except Exception as e:
        import traceback

        print(" Regeneration Error:")
        print(traceback.format_exc())

        raise HTTPException(
            status_code=500, detail=f"Failed to regenerate framework: {str(e)}"
        )


def convert_framework_to_text(framework: dict) -> str:
    """
    Convert the frame JSON back to text (for local LLM processing)
    This is a simplified version used to simulate the original document.
    """
    parts = []

    # Title
    metadata = framework.get("metadata", {})
    title = metadata.get("title", "Framework")
    parts.append(f"# {title}\n")

    # Steps
    steps = framework.get("steps", [])
    if steps:
        parts.append("\n## Framework Steps\n")
        for step in steps:
            parts.append(f"\n### {step.get('name', 'Step')}")
            parts.append(step.get("description", ""))
            sub_steps = step.get("subSteps", [])
            if sub_steps:
                for sub in sub_steps:
                    parts.append(f"- {sub}")

    # Risks
    risks = framework.get("risks", [])
    if risks:
        parts.append("\n## Risks\n")
        for risk in risks:
            parts.append(f"\n### {risk.get('title', 'Risk')}")
            parts.append(risk.get("description", ""))

    # Escalation
    escalation = framework.get("escalation", [])
    if escalation:
        parts.append("\n## Escalation Points\n")
        for esc in escalation:
            parts.append(f"- When: {esc.get('trigger', 'Unknown')}")
            parts.append(f"  Action: {esc.get('action', 'Escalate')}")

    return "\n".join(parts)


# ============= AI Merge Endpoint =============


class AIMergeRequest(BaseModel):
    """AI merge request model"""

    frameworks: List[dict]  # Multiple frameworks selected by the user


@router.post("/ai-merge")
async def ai_merge_frameworks(
    request: AIMergeRequest, current_user_id: str = Depends(get_current_user_id)
):
    """
    Use AI to intelligently merge multiple frameworks.

    Requires a valid JWT before any merge or mock-merge logic can run.
    """
    try:
        # Validate input
        if not request.frameworks or len(request.frameworks) < 2:
            raise HTTPException(
                status_code=400, detail="Please select at least 2 frameworks to merge"
            )

        if len(request.frameworks) > 10:
            raise HTTPException(
                status_code=400, detail="Cannot merge more than 10 frameworks at once"
            )

        print(f"🔀 AI Merge: Merging {len(request.frameworks)} frameworks")

        # check API key
        api_key, base_url = resolve_api_settings(None, None)
        if not api_key:
            # If no API key is specified, a simple merge result is returned.
            print("⚠️ No API key, returning mock merge")
            return {
                "success": True,
                "merged_framework": {
                    "name": "Merged Framework (Mock)",
                    "description": "This is a test merge of "
                    + str(len(request.frameworks))
                    + " frameworks.",
                    "subSteps": [
                        "Combined step 1",
                        "Combined step 2",
                        "Combined step 3",
                    ],
                },
            }

        # Ready to merge prompt
        frameworks_text = []
        for i, fw in enumerate(request.frameworks, 1):
            frameworks_text.append(f"\n{'='*60}")
            frameworks_text.append(f"FRAMEWORK {i}: {fw.get('name', 'Unnamed')}")
            frameworks_text.append(f"{'='*60}\n")

            # Description
            if fw.get("description"):
                frameworks_text.append(f"Description:\n{fw['description']}\n")

            # Sub-steps
            if fw.get("subSteps"):
                frameworks_text.append("Sub-steps:")
                for j, step in enumerate(fw["subSteps"], 1):
                    frameworks_text.append(f"  {j}. {step}")
                frameworks_text.append("")

        combined_text = "\n".join(frameworks_text)

        # Call OpenAI
        from openai import OpenAI
        import httpx
        import os

        # Clear proxy environment variables
        original_proxies = {}
        proxy_keys = [
            "HTTP_PROXY",
            "HTTPS_PROXY",
            "http_proxy",
            "https_proxy",
            "ALL_PROXY",
            "all_proxy",
            "NO_PROXY",
            "no_proxy",
        ]

        for key in proxy_keys:
            if key in os.environ:
                original_proxies[key] = os.environ[key]
                del os.environ[key]

        try:
            # OpenAI 2.6.1 automatically handles retries and timeouts.
            if base_url:
                client = OpenAI(
                    api_key=api_key,
                    base_url=base_url,
                    timeout=300.0,  # 5 minutes timeout
                    max_retries=2,
                )
            else:
                client = OpenAI(api_key=api_key, timeout=300.0, max_retries=2)

            # Build prompt
            system_prompt = (
                "You are a framework merging assistant. "
                "Your task is to intelligently combine multiple frameworks into one cohesive framework. "
                "You should:\n"
                "1. Identify common themes and consolidate similar content\n"
                "2. Remove redundancy while preserving unique insights from each framework\n"
                "3. Organize the merged content logically\n"
                "4. Create a clear, comprehensive description that captures all key aspects\n"
                "5. Combine sub-steps in a logical order\n"
                "6. Generate an appropriate name for the merged framework\n\n"
                "Return ONLY a valid JSON object with this structure:\n"
                "{\n"
                '  "name": "Merged Framework Name",\n'
                '  "description": "Comprehensive description...",\n'
                '  "subSteps": ["Step 1", "Step 2", ...]\n'
                "}"
            )

            user_prompt = (
                f"Please merge these {len(request.frameworks)} frameworks into one:\n\n"
                f"{combined_text}\n\n"
                "Create a new framework that:\n"
                "- Captures the essence of all input frameworks\n"
                "- Eliminates redundancy and contradictions\n"
                "- Provides a clear, actionable structure\n"
                "- Has a descriptive name that reflects the merged content\n\n"
                "Return the merged framework as JSON."
            )

            print(" Sending merge request to OpenAI...")
            response = client.chat.completions.create(
                model="gpt-4o",
                temperature=0.4,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
            )

            result_text = response.choices[0].message.content.strip()
            print(" Received response from OpenAI")

            # Parsing JSON
            from llm_global import robust_json_loads

            merged_framework = robust_json_loads(result_text)

            # Ensure required fields exist
            if not merged_framework.get("name"):
                merged_framework["name"] = "AI Merged Framework"

            if not merged_framework.get("description"):
                merged_framework["description"] = ""

            if not merged_framework.get("subSteps"):
                merged_framework["subSteps"] = []

            print(f" Successfully merged into: {merged_framework['name']}")

            return {"success": True, "merged_framework": merged_framework}

        finally:
            # Restore proxy settings
            for key, value in original_proxies.items():
                os.environ[key] = value

            # Close HTTP client
            if "http_client" in locals():
                try:
                    http_client.close()
                except:
                    pass

    except HTTPException:
        raise
    except Exception as e:
        import traceback

        print(" AI Merge Error:")
        print(traceback.format_exc())

        return {"success": False, "error": str(e)}


# ============= AI Fill Endpoint =============


class AIFillRequest(BaseModel):
    """AI Fill request model"""

    artefact_name: str
    artefact_summary: str = ""
    existing_sections: List[dict] = []  # [{heading: str, body: str}]
    sections_to_fill: List[str] = []  # List of section headings to fill


def parse_ai_json_response(text):
    """
    Parse JSON from AI response, handling common issues like markdown code blocks
    """
    import json
    import re

    # Remove markdown code blocks if present
    text = text.strip()

    # Remove ```json ... ``` or ``` ... ```
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    text = text.strip()

    # Try to find JSON array in the text
    # Look for [ ... ] pattern
    match = re.search(r"\[[\s\S]*\]", text)
    if match:
        text = match.group(0)

    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        print(f"JSON parse error: {e}")
        print(f"Attempted to parse: {text[:500]}...")
        raise


@router.post("/ai-fill")
async def ai_fill_sections(
    request: AIFillRequest, current_user_id: str = Depends(get_current_user_id)
):
    """
    Use AI to fill empty section content based on context
    """
    try:
        if not request.sections_to_fill:
            return {
                "success": True,
                "filled_sections": [],
                "message": "No sections to fill",
            }

        print(
            f"✨ AI Fill: Filling {len(request.sections_to_fill)} sections for '{request.artefact_name}'"
        )

        # Check API key
        api_key, base_url = resolve_api_settings(None, None)
        if not api_key:
            # Mock response if no API key
            print("⚠️ No API key, returning mock fill")
            filled = []
            for section_name in request.sections_to_fill:
                filled.append(
                    {
                        "heading": section_name,
                        "body": f"[Mock content for {section_name}] This is placeholder content. Configure your OpenAI API key for real AI-generated content.",
                    }
                )
            return {"success": True, "filled_sections": filled}

        # Build context from existing sections
        existing_context = ""
        if request.existing_sections:
            existing_context = "\n\nExisting sections for context:\n"
            for sec in request.existing_sections:
                existing_context += (
                    f"- {sec.get('heading', 'Section')}: {sec.get('body', '')}\n"
                )

        # Prepare OpenAI call
        from openai import OpenAI
        import os

        # Clear proxy settings
        original_proxies = {}
        proxy_keys = [
            "HTTP_PROXY",
            "HTTPS_PROXY",
            "http_proxy",
            "https_proxy",
            "ALL_PROXY",
            "all_proxy",
        ]
        for key in proxy_keys:
            if key in os.environ:
                original_proxies[key] = os.environ[key]
                del os.environ[key]

        try:
            if base_url:
                client = OpenAI(
                    api_key=api_key, base_url=base_url, timeout=120.0, max_retries=2
                )
            else:
                client = OpenAI(api_key=api_key, timeout=120.0, max_retries=2)

            # Build prompt
            system_prompt = (
                "You are an expert document writer. Your task is to fill in content for document sections. "
                "If section names are numbers or very short, infer logical section topics based on the document context. "
                "You MUST return ONLY a valid JSON array with no additional text, markdown, or explanation. "
                "Do not wrap the response in code blocks. Just output the raw JSON array."
            )

            sections_list = ", ".join([f'"{s}"' for s in request.sections_to_fill])

            # Check if sections are just numbers - provide extra context
            has_number_sections = any(
                s.strip().isdigit() for s in request.sections_to_fill
            )
            extra_instruction = ""
            if has_number_sections:
                extra_instruction = """
IMPORTANT: Some sections are numbered (e.g., "3", "4"). Based on the document type and existing sections,
infer what these numbered sections should logically contain. For example:
- For compliance documents: risk analysis, mitigation steps, monitoring procedures
- For technical documents: implementation details, testing procedures, maintenance
- Follow the pattern of existing numbered sections if present."""

            user_prompt = f"""Document: {request.artefact_name}
Summary: {request.artefact_summary or 'A professional document'}
{existing_context}
{extra_instruction}

Write professional content for these sections: {sections_list}

Output format (return ONLY this JSON array, nothing else):
[{{"heading": "Section Name", "body": "Professional content here..."}}]

Requirements:
- Each body should be 2-3 sentences of relevant, professional content
- Infer appropriate content based on document type and context
- Be specific and actionable, not generic"""

            print("📡 Sending request to OpenAI...")
            response = client.chat.completions.create(
                model="gpt-4o",
                temperature=0.4,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
            )

            result_text = response.choices[0].message.content.strip()
            print(f"✅ Received response from OpenAI: {result_text[:200]}...")

            # Parse JSON response with our custom parser
            filled_sections = parse_ai_json_response(result_text)

            # Validate response format
            if not isinstance(filled_sections, list):
                filled_sections = [filled_sections]

            # Ensure all requested sections are filled
            filled_headings = {s.get("heading") for s in filled_sections}
            for section_name in request.sections_to_fill:
                if section_name not in filled_headings:
                    filled_sections.append(
                        {
                            "heading": section_name,
                            "body": f"Content for {section_name} section.",
                        }
                    )

            print(f"✅ Successfully filled {len(filled_sections)} sections")
            return {"success": True, "filled_sections": filled_sections}

        finally:
            # Restore proxy settings
            for key, value in original_proxies.items():
                os.environ[key] = value

    except Exception as e:
        import traceback

        print("❌ AI Fill Error:")
        print(traceback.format_exc())
        return {"success": False, "error": str(e), "filled_sections": []}


@router.post("/sync-library")
def sync_library(
    req: SyncLibraryRequest,
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db),
):
    import requests as _requests

    project_id = _read_project_id(req.project_id)
    api_key = _read_api_key(req.api_key)
    id_token = _ensure_id_token(api_key, req.id_token)
    vs_id = req.vector_store_id or os.getenv("OPENAI_VECTOR_STORE_LIBRARY")
    if not project_id or not api_key or not id_token or not vs_id:
        raise HTTPException(status_code=400, detail="missing parameters")

    client = _openai_client(use_vector_store_key=True)
    vs_api = _vs_api(client)
    url = f"https://firestore.googleapis.com/v1/projects/{project_id}/databases/(default)/documents:runQuery"
    headers = {"Authorization": f"Bearer {id_token}"}
    params = {"key": api_key}
    body_public = {
        "structuredQuery": {
            "from": [{"collectionId": "frameworks"}],
            "where": {
                "fieldFilter": {
                    "field": {"fieldPath": "isPublic"},
                    "op": "EQUAL",
                    "value": {"booleanValue": True},
                }
            },
            "limit": req.limit,
        }
    }
    r = _requests.post(
        url, headers=headers, params=params, json=body_public, timeout=60
    )
    if r.status_code != 200:
        raise HTTPException(status_code=500, detail=r.text)
    items = {}
    for it in r.json():
        doc = it.get("document")
        if not doc:
            continue
        name = doc.get("name", "")
        doc_id = name.split("/")[-1] if name else generate()
        fields = doc.get("fields", {})
        parsed = {k: _val(fields[k]) for k in fields}
        parsed["id"] = doc_id
        items[doc_id] = parsed

    if req.include_organization:
        body_org = {
            "structuredQuery": {
                "from": [{"collectionId": "frameworks"}],
                "where": {
                    "fieldFilter": {
                        "field": {"fieldPath": "publishedToOrganization"},
                        "op": "EQUAL",
                        "value": {"booleanValue": True},
                    }
                },
                "limit": req.limit,
            }
        }
        r2 = _requests.post(
            url, headers=headers, params=params, json=body_org, timeout=60
        )
        if r2.status_code != 200:
            raise HTTPException(status_code=500, detail=r2.text)
        for it in r2.json():
            doc = it.get("document")
            if not doc:
                continue
            name = doc.get("name", "")
            doc_id = name.split("/")[-1] if name else generate()
            fields = doc.get("fields", {})
            parsed = {k: _val(fields[k]) for k in fields}
            parsed["id"] = doc_id
            items[doc_id] = parsed

    uploaded = 0
    for doc_id, obj in items.items():
        exists = (
            db.query(SyncedVectorItem)
            .filter(SyncedVectorItem.doc_id == doc_id, SyncedVectorItem.vs_id == vs_id)
            .first()
        )
        if exists:
            continue

        fn = f"framework_{doc_id}.json"

        # get attributes
        is_public = obj.get("isPublic", False)
        tenant_id = obj.get("tenantId", "")

        attrs = {
            "visibility": "public" if is_public else "private",
            "domain": "valorie.ai",
            "tenantId": tenant_id,
        }

        # upload Vector Store
        _vs_upload_json(
            client,
            vs_api,
            vs_id,
            fn,
            json.dumps(obj, ensure_ascii=False),
            attributes=attrs,
        )

        # save record
        rec = SyncedVectorItem(
            doc_id=doc_id,
            vs_id=vs_id,
            source=("organization" if obj.get("publishedToOrganization") else "public"),
        )
        db.add(rec)
        try:
            db.commit()
        except Exception:
            db.rollback()
        uploaded += 1
    return {"success": True, "uploaded": uploaded}


@router.post("/log-event")
def log_event(
    req: EventLogRequest, user_id: str = Depends(get_current_user_id)
):
    vs_id = os.getenv("OPENAI_VECTOR_STORE_ACTIVITY")
    if not vs_id:
        raise HTTPException(
            status_code=500, detail="OPENAI_VECTOR_STORE_ACTIVITY missing"
        )
    client = _openai_client(use_vector_store_key=True)
    vs_api = _vs_api(client)
    now = datetime.utcnow().isoformat()
    payload = {
        "type": req.type,
        "frameworkId": req.framework_id,
        "tenantId": req.tenant_id,
        "userId": user_id,
        "timestamp": now,
        "data": req.payload or {},
    }
    fn = f"event_{generate()}.json"
    _vs_upload_json(client, vs_api, vs_id, fn, json.dumps(payload, ensure_ascii=False))
    return {"success": True}


@router.post("/push-framework")
def push_framework(
    req: PushFrameworkRequest, current_user_id: str = Depends(get_current_user_id)
):
    print("=" * 60)
    print("🔧 DEBUG: push_framework called")
    print(f"   - Request data: {req}")
    print(f"   - Framework ID: {req.framework.get('id') if req.framework else 'None'}")
    print(f"   - Vector Store ID (req): {req.vector_store_id}")
    print(
        f"   - OPENAI_VECTOR_STORE_LIBRARY: {os.getenv('OPENAI_VECTOR_STORE_LIBRARY')}"
    )
    print(
        f"   - OPENAI_VECTOR_STORE_ACTIVITY: {os.getenv('OPENAI_VECTOR_STORE_ACTIVITY')}"
    )
    print(
        f"   - OPENAI_VECTOR_STORE_API_KEY: {'Set' if os.getenv('OPENAI_VECTOR_STORE_API_KEY') else 'Missing'}"
    )
    print("=" * 60)

    try:
        client = _openai_client(use_vector_store_key=True)
        vs_api = _vs_api(client)
        vs_id = (
            req.vector_store_id
            or os.getenv("OPENAI_VECTOR_STORE_LIBRARY")
            or os.getenv("OPENAI_VECTOR_STORE_ACTIVITY")
        )

        print(f"🔧 DEBUG: Final vs_id = {vs_id}")

        if not vs_id:
            print("❌ ERROR: vs_id is None!")
            raise HTTPException(status_code=400, detail="vector_store_id missing")

        fw = req.framework or {}
        doc_id = fw.get("id") or generate()
        fn = f"framework_{doc_id}.json"

        is_public = fw.get("isPublic", False)
        tenant_id = fw.get("tenantId", "")

        print(f"🔧 DEBUG: Framework metadata:")
        print(f"   - doc_id: {doc_id}")
        print(f"   - filename: {fn}")
        print(f"   - isPublic: {is_public}")
        print(f"   - tenantId: {tenant_id}")

        attrs = {
            "visibility": "public" if is_public else "private",
            "domain": "valorie.ai",
            "tenantId": tenant_id,
        }

        print(f"🔧 DEBUG: Uploading to Vector Store...")
        _vs_upload_json(
            client,
            vs_api,
            vs_id,
            fn,
            json.dumps(fw, ensure_ascii=False),
            attributes=attrs,
        )
        print(f"✅ DEBUG: Upload successful!")

        return {"success": True}

    except HTTPException as he:
        print(f"❌ ERROR (HTTPException): {he.detail}")
        raise
    except Exception as e:
        print(f"❌ ERROR (Exception): {str(e)}")
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
